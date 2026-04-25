"""
Worker Report — Tarea Celery para generar informes de escaneo.

Recibe scan_id, dominio y ruta del CSV, analiza los datos
y genera informe en PDF usando ReportLab.
Registra el informe en la base de datos.
"""

import os
import re
import math
import json
import logging
from datetime import datetime, timezone
from io import BytesIO
from celery import Celery
from sqlalchemy import create_engine, text

logger = logging.getLogger(__name__)

CELERY_BROKER_URL = os.getenv("CELERY_BROKER_URL", "amqp://asm:password@rabbitmq:5672//")
REPORTS_DIR = os.getenv("REPORTS_DIR", "/data/reportes")
CONSOLIDATED_DIR = os.getenv("CONSOLIDATED_DIR", "/data/consolidados")

DATABASE_URL = (
    f"postgresql://{os.getenv('POSTGRES_USER', 'asm_user')}:"
    f"{os.getenv('POSTGRES_PASSWORD')}@"
    f"{os.getenv('POSTGRES_HOST', 'db')}:"
    f"{os.getenv('POSTGRES_PORT', '5432')}/"
    f"{os.getenv('POSTGRES_DB', 'asm_db')}"
)

app = Celery("worker_report", broker=CELERY_BROKER_URL)
engine = create_engine(DATABASE_URL, pool_pre_ping=True)

# ── Constantes de análisis (portadas del monolito) ─────────────────
PUERTOS_RIESGO = {
    21: ("FTP", "Alto"), 22: ("SSH", "Alto"), 23: ("Telnet", "Crítico"),
    25: ("SMTP", "Alto"), 80: ("HTTP", "Bajo"), 443: ("HTTPS", "Bajo"),
    3306: ("MySQL", "Alto"), 3389: ("RDP", "Crítico"), 5432: ("PostgreSQL", "Alto"),
    5900: ("VNC", "Crítico"), 6379: ("Redis", "Alto"), 8080: ("HTTP Proxy", "Alto"),
    8443: ("HTTPS Alt", "Bajo"), 27017: ("MongoDB", "Alto"),
}

SEVERITY_WEIGHT = {"Crítica": 4, "Alta": 3, "Media": 2, "Baja": 1, "Informativa": 0}

CIFRADOS_DEBILES_REGEX = r"(?:RC4|CBC|DES|NULL|EXPORT|MD5)"
TLS_OBSOLETO_REGEX = r"tls1\.0|tls1_0|tls1\.1|tls1_1|sslv"


def register_report(scan_id: int, fmt: str, path: str):
    with engine.connect() as conn:
        conn.execute(
            text("INSERT INTO reports (scan_id, format, path) VALUES (:sid, :fmt, :path)"),
            {"sid": scan_id, "fmt": fmt, "path": path},
        )
        conn.commit()


def evaluar_fila(row: dict) -> dict:
    """Evalúa una fila del CSV y devuelve flags de vulnerabilidad."""
    flags = {}
    puertos_str = str(row.get("Puertos abiertos", ""))
    puertos = [p.strip() for p in puertos_str.split(",") if p.strip().isdigit()]

    flags["flag_subdominios_huerfanos"] = str(row.get("Estado", "")).lower() in ("huérfano", "huerfano", "")
    flags["flag_carencia_spf_dkim_dmarc"] = "ausente" in str(row.get("Correo", "")).lower() or \
                                             "sin dmarc" in str(row.get("Correo", "")).lower()
    flags["flag_software_expuesto"] = str(row.get("version servidor", "")).lower() not in (
        "", "no identificado", "no expuesto"
    )
    flags["flag_exposicion_puertos_admin"] = any(
        p in ["22", "23", "3389", "5900", "5432", "3306", "8080", "10000"] for p in puertos
    )
    flags["flag_headers_esenciales"] = "ausente" in str(row.get("Estado Certificado", "")).lower()
    flags["flag_cert_ssl_invalido"] = str(row.get("Estado Certificado", "")).lower() in (
        "vencido", "inválido", "invalido", "sin ssl"
    )
    tls = str(row.get("Versiones TLS soportadas", "")).lower()
    flags["flag_tls_obsoleto"] = bool(re.search(TLS_OBSOLETO_REGEX, tls))
    cifrados = str(row.get("Riesgos de cifrado", "")).lower()
    flags["flag_cifrados_debiles"] = bool(re.search(CIFRADOS_DEBILES_REGEX, cifrados, re.IGNORECASE))
    return flags


def calcular_criticidad(df) -> tuple:
    score = 0
    flag_cols = [c for c in df.columns if c.startswith("flag_")]
    for _, row in df.iterrows():
        for col in flag_cols:
            if row.get(col):
                score += 2
    if score >= 24: return "Crítica", score
    elif score >= 14: return "Alta", score
    elif score >= 6: return "Media", score
    elif score >= 1: return "Baja", score
    return "Informativa", score


def generar_pdf(scan_id: int, domain: str, df, nivel: str, score: int) -> str:
    """Genera informe PDF con ReportLab."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.units import cm
    except ImportError:
        logger.error("ReportLab no disponible")
        return ""

    os.makedirs(os.path.join(REPORTS_DIR, "pdf"), exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_domain = domain.replace(".", "_")
    pdf_path = os.path.join(REPORTS_DIR, "pdf", f"informe_{safe_domain}_{ts}.pdf")

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = []

    # Título
    story.append(Paragraph(f"Informe de Superficie de Ataque", styles["Title"]))
    story.append(Paragraph(f"Dominio: <b>{domain}</b>", styles["Heading2"]))
    story.append(Paragraph(f"Criticidad global: <b>{nivel}</b> (score: {score})", styles["Normal"]))
    story.append(Paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    # Resumen de hallazgos
    story.append(Paragraph("Resumen de Hallazgos", styles["Heading2"]))
    flag_cols = [c for c in df.columns if c.startswith("flag_")]
    hallazgos_data = [["Categoría", "Afectados"]]
    nombres = {
        "flag_subdominios_huerfanos": "Subdominios huérfanos",
        "flag_carencia_spf_dkim_dmarc": "Carencia SPF/DKIM/DMARC",
        "flag_software_expuesto": "Software expuesto",
        "flag_exposicion_puertos_admin": "Puertos de administración expuestos",
        "flag_headers_esenciales": "Cabeceras HTTP ausentes",
        "flag_cert_ssl_invalido": "Certificado SSL inválido",
        "flag_tls_obsoleto": "Protocolos TLS obsoletos",
        "flag_cifrados_debiles": "Cifrados débiles",
    }
    for col in flag_cols:
        total = int(df[col].sum())
        if total > 0:
            hallazgos_data.append([nombres.get(col, col), str(total)])

    if len(hallazgos_data) > 1:
        t = Table(hallazgos_data, colWidths=[12*cm, 3*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1d4ed8")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f1f5f9")]),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ]))
        story.append(t)
    else:
        story.append(Paragraph("No se detectaron vulnerabilidades significativas.", styles["Normal"]))

    story.append(Spacer(1, 0.5*cm))

    # Detalle de activos
    story.append(Paragraph("Detalle de Activos Analizados", styles["Heading2"]))
    cols_show = ["Subdominio", "Estado", "Puertos abiertos", "Estado Certificado"]
    cols_exist = [c for c in cols_show if c in df.columns]
    table_data = [cols_exist] + df[cols_exist].fillna("").values.tolist()
    if len(table_data) > 1:
        t2 = Table(table_data[:51])  # máx 50 filas
        t2.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e2e8f0")),
        ]))
        story.append(t2)

    doc.build(story)

    with open(pdf_path, "wb") as f:
        f.write(buffer.getvalue())

    return pdf_path


def generar_docx(scan_id: int, domain: str, df, nivel: str, score: int) -> str:
    """Genera informe DOCX con python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        logger.error("python-docx no disponible")
        return ""

    os.makedirs(os.path.join(REPORTS_DIR, "docx"), exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_domain = domain.replace(".", "_")
    docx_path = os.path.join(REPORTS_DIR, "docx", f"informe_{safe_domain}_{ts}.docx")

    doc = Document()

    # Estilos base
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Portada
    titulo = doc.add_heading('Informe de Superficie de Ataque', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(f'Dominio: {domain}', level=1)
    doc.add_paragraph(f'Criticidad global: {nivel}  |  Score: {score}')
    doc.add_paragraph(f'Fecha de generación: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph(f'Escaneo ID: {scan_id}')
    doc.add_paragraph('')

    # ── Resumen de hallazgos ──────────────────────────────
    doc.add_heading('Resumen de Hallazgos', level=1)

    flag_cols = [c for c in df.columns if c.startswith("flag_")]
    nombres = {
        "flag_subdominios_huerfanos":      "Subdominios huérfanos",
        "flag_carencia_spf_dkim_dmarc":    "Carencia SPF/DKIM/DMARC",
        "flag_software_expuesto":           "Software expuesto",
        "flag_exposicion_puertos_admin":    "Puertos de administración expuestos",
        "flag_headers_esenciales":          "Cabeceras HTTP ausentes",
        "flag_cert_ssl_invalido":           "Certificado SSL inválido",
        "flag_tls_obsoleto":                "Protocolos TLS obsoletos",
        "flag_cifrados_debiles":            "Cifrados débiles",
    }
    severidades = {
        "flag_subdominios_huerfanos":      "Media",
        "flag_carencia_spf_dkim_dmarc":    "Alta",
        "flag_software_expuesto":           "Alta",
        "flag_exposicion_puertos_admin":    "Alta",
        "flag_headers_esenciales":          "Media",
        "flag_cert_ssl_invalido":           "Alta",
        "flag_tls_obsoleto":                "Media",
        "flag_cifrados_debiles":            "Media",
    }

    hallazgos_data = [(c, int(df[c].sum()), severidades.get(c, "Media")) for c in flag_cols if c in df.columns and df[c].sum() > 0]

    if hallazgos_data:
        tabla = doc.add_table(rows=1, cols=3)
        tabla.style = 'Table Grid'
        hdr = tabla.rows[0].cells
        hdr[0].text = 'Categoría'
        hdr[1].text = 'Afectados'
        hdr[2].text = 'Severidad'
        for cell in hdr:
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
            run.bold = True

        for col, total, sev in hallazgos_data:
            row = tabla.add_row().cells
            row[0].text = nombres.get(col, col)
            row[1].text = str(total)
            row[2].text = sev
    else:
        doc.add_paragraph('No se detectaron vulnerabilidades significativas.')

    doc.add_paragraph('')

    # ── Detalle de activos ────────────────────────────────
    doc.add_heading('Detalle de Activos Analizados', level=1)

    cols_show = ["Subdominio", "Estado", "Puertos abiertos", "Estado Certificado", "Correo", "version servidor"]
    cols_exist = [c for c in cols_show if c in df.columns]

    if cols_exist and not df.empty:
        tabla2 = doc.add_table(rows=1, cols=len(cols_exist))
        tabla2.style = 'Table Grid'
        hdr2 = tabla2.rows[0].cells
        for i, col in enumerate(cols_exist):
            hdr2[i].text = col
            if hdr2[i].paragraphs[0].runs:
                hdr2[i].paragraphs[0].runs[0].bold = True

        for _, row in df[cols_exist].fillna("").head(100).iterrows():
            cells = tabla2.add_row().cells
            for i, col in enumerate(cols_exist):
                cells[i].text = str(row[col])[:80]  # truncar celdas largas

    doc.save(docx_path)
    return docx_path


def actualizar_consolidado(domain: str, df):
    """Actualiza el JSON consolidado multientidad."""
    os.makedirs(CONSOLIDATED_DIR, exist_ok=True)
    path = os.path.join(CONSOLIDATED_DIR, "consolidado_data.json")

    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = {"entities": [], "last_updated": ""}

    # Eliminar entrada previa del mismo dominio
    data["entities"] = [e for e in data["entities"] if e.get("domain") != domain]

    # Agregar nueva entrada
    flag_cols = [c for c in df.columns if c.startswith("flag_")]
    data["entities"].append({
        "domain": domain,
        "total_assets": len(df),
        "active_assets": int((df.get("Estado", "").str.lower() == "activo").sum()) if "Estado" in df.columns else 0,
        "flags": {c: int(df[c].sum()) for c in flag_cols if c in df.columns},
        "scan_date": datetime.now(timezone.utc).isoformat(),
    })
    data["last_updated"] = datetime.now(timezone.utc).isoformat()

    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


@app.task(name="worker_report.tasks.generate_report", bind=True, max_retries=0)
def generate_report(self, scan_id: int, domain: str, csv_path: str):
    """
    Genera informe PDF para un escaneo completado.
    Registra el informe en la base de datos.
    Actualiza el JSON consolidado.
    """
    import pandas as pd

    logger.info(f"[report:{scan_id}] Generando informe para {domain}")

    if not os.path.exists(csv_path):
        logger.error(f"[report:{scan_id}] CSV no encontrado: {csv_path}")
        return

    try:
        df = pd.read_csv(csv_path).fillna("")

        # Evaluar flags por fila
        flag_rows = [evaluar_fila(row) for _, row in df.iterrows()]
        flags_df = pd.DataFrame(flag_rows)
        df = pd.concat([df, flags_df], axis=1)

        nivel, score = calcular_criticidad(df)
        logger.info(f"[report:{scan_id}] Criticidad: {nivel} (score={score})")

        # Generar PDF
        pdf_path = generar_pdf(scan_id, domain, df, nivel, score)
        if pdf_path:
            register_report(scan_id, "pdf", pdf_path)
            logger.info(f"[report:{scan_id}] PDF generado: {pdf_path}")

        # Generar DOCX
        docx_path = generar_docx(scan_id, domain, df, nivel, score)
        if docx_path:
            register_report(scan_id, "docx", docx_path)
            logger.info(f"[report:{scan_id}] DOCX generado: {docx_path}")

        # Actualizar consolidado
        actualizar_consolidado(domain, df)
        logger.info(f"[report:{scan_id}] Consolidado actualizado")

    except Exception as e:
        logger.error(f"[report:{scan_id}] Error: {e}", exc_info=True)
